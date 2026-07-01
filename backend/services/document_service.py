# -*- coding: utf-8 -*-
"""
=============================================================================
文档服务模块
=============================================================================

功能:
- 解析 Word 文档 (.docx) 和纯文本 (.txt)
- 语义分块 (Semantic Chunking)
- 向量化存储到 ChromaDB
- 两阶段检索 + LLM 整合
- 问答机器人

=============================================================================
"""

import uuid
import io
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument
from langchain_openai import ChatOpenAI
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
import chromadb
from chromadb.config import Settings
import os

from config import LLM_API_KEY, LLM_API_BASE, LLM_MODEL

# 获取 API Key
_deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
if not _deepseek_api_key:
    raise ValueError("请设置 DEEPSEEK_API_KEY 环境变量")


# ============================================================================
# ChromaDB 客户端管理
# ============================================================================

_chroma_client: Optional[Chroma] = None


def get_chroma_client() -> Chroma:
    """获取 Chroma 向量库客户端（单例）"""
    global _chroma_client
    if _chroma_client is None:
        persist_dir = os.environ.get("CHROMA_PERSIST_DIR", "/app/chroma_data")
        os.makedirs(persist_dir, exist_ok=True)
        embeddings = get_embeddings()
        ## 数据持续到磁盘而非内存
        client = chromadb.PersistentClient(
            path=persist_dir,
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True,
            ),
        )
        _chroma_client = Chroma(
            client=client,
            collection_name="documents",
            embedding_function=embeddings,
        )
    return _chroma_client


# ============================================================================
# Embedding 配置
# ============================================================================


def get_embeddings():
    """获取 Embedding 函数（本地 BAAI/bge-small-zh-v1.5，中文语义检索）"""
    model_name = os.environ.get(
        "EMBEDDING_MODEL",
        "BAAI/bge-small-zh-v1.5",
    )
    cache_dir = os.environ.get("EMBEDDING_CACHE_DIR", "/app/models")
    os.makedirs(cache_dir, exist_ok=True)
    return HuggingFaceEmbeddings(
        model_name=model_name,
        cache_folder=cache_dir,
        model_kwargs={"device": "cpu"},  # 使用CPU进行推理
        encode_kwargs={
            "normalize_embeddings": True
        },  # L2 归一化，使余弦相似度等价于内积，提升检索效率
    )


# ============================================================================
# 文档解析（支持 docx 和 txt）
# ============================================================================


def parse_txt_document(file_content: bytes) -> Tuple[str, List[Dict[str, Any]]]:
    """
    解析 txt 纯文本文件

    Args:
        file_content: 文件二进制内容

    Returns:
        (full_text, content_items):
            - full_text: 完整文本内容
            - content_items: 按段落拆分的列表，每项包含 text 和 metadata
    """
    text = file_content.decode("utf-8")
    lines = text.split("\n")
    content_items = []
    current_para = []
    current_headings = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_para:
                content_items.append(
                    {
                        "text": "\n".join(current_para),
                        "type": "paragraph",
                        "headings": list(current_headings),
                    }
                )
                current_para = []
            continue

        # 检测标题行（以 # 开头 或 全大写 或 短行+冒号结尾）
        is_heading = (
            stripped.startswith("#")
            or (len(stripped) < 50 and stripped.isupper())
            or (
                len(stripped) < 60
                and stripped.endswith(":")
                and not stripped.endswith("::")
            )
        )

        if is_heading:
            if current_para:
                content_items.append(
                    {
                        "text": "\n".join(current_para),
                        "type": "paragraph",
                        "headings": list(current_headings),
                    }
                )
                current_para = []
            content_items.append(
                {
                    "text": stripped,
                    "type": "heading",
                    "level": 1 if stripped.startswith("#") else 2,
                    "headings": list(current_headings),
                }
            )
            if not stripped.startswith("#"):
                current_headings.append(stripped)
        else:
            current_para.append(stripped)

    # 处理最后一段
    if current_para:
        content_items.append(
            {
                "text": "\n".join(current_para),
                "type": "paragraph",
                "headings": list(current_headings),
            }
        )

    full_text = "\n\n".join(item["text"] for item in content_items)
    return full_text, content_items


def extract_heading_chain(
    paragraphs: List[Paragraph],
) -> Dict[int, Tuple[int, str, str]]:
    """
    提取文档中的标题结构

    Returns:
        Dict[paragraph_index, (level, heading_text, full_chain)]
        例如: {3: (1, "第一章", "第一章"), 7: (2, "第一节", "第一章 > 第一节")}
    """
    heading_info = {}
    current_chain: List[str] = []

    for idx, para in enumerate(paragraphs):
        if para.style.name.startswith("Heading"):
            # 提取标题级别
            level = (
                int(para.style.name.replace("Heading ", ""))
                if "Heading " in para.style.name
                else 1
            )
            text = para.text.strip()

            # 更新标题链
            current_chain = current_chain[: level - 1] + [text]
            full_chain = " > ".join(current_chain)

            heading_info[idx] = (level, text, full_chain)

    return heading_info


def parse_word_document(file_content: bytes) -> Tuple[str, List[Tuple[str, str, int]]]:
    """
    解析 Word 文档，提取文本和标题信息

    Returns:
        (full_text, [(paragraph_text, heading_chain, para_index), ...])
    """
    doc = Document(io.BytesIO(file_content))
    paragraphs = list(doc.paragraphs)
    heading_info = extract_heading_chain(paragraphs)

    content_items = []
    full_text_parts = []

    for idx, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue

        full_text_parts.append(text)

        # 获取该段落的标题链
        heading_chain = ""
        for i in range(idx, -1, -1):
            if i in heading_info:
                heading_chain = heading_info[i][2]
                break

        content_items.append((text, heading_chain, idx))

    # 提取表格内容
    table_items, table_text = extract_tables(doc, len(paragraphs), heading_info)
    content_items.extend(table_items)

    if table_text:
        full_text_parts.append(table_text)

    full_text = "\n\n".join(full_text_parts)
    return full_text, content_items


def extract_tables(
    doc: Document, base_index: int, heading_info: Dict[int, Tuple[int, str, str]]
) -> Tuple[List[Tuple[str, str, int]], str]:
    """
    从 Word 文档中提取表格内容

    Args:
        doc: Document 对象
        base_index: 基础段落索引，用于计算表格索引
        heading_info: 标题信息字典

    Returns:
        (table_items, table_text): 表格项列表和合并的表格文本
    """
    table_items = []
    table_text_parts = []

    for table_idx, table in enumerate(doc.tables):
        table_index = base_index + table_idx
        table_text = format_table_as_text(table)

        if table_text.strip():
            table_items.append((table_text, "[表格]", table_index))
            table_text_parts.append(table_text)

    table_text = (
        "\n\n[表格内容]\n" + "\n".join(table_text_parts) if table_text_parts else ""
    )
    return table_items, table_text


def format_table_as_text(table) -> str:
    """
    将表格格式化为易读的文本格式

    使用分隔符将表格行列转换为文本，保留表格结构信息。
    """
    if not table.rows:
        return ""

    rows_text = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        row_text = " | ".join(cells)
        if row_text.strip():
            rows_text.append(row_text)

    if not rows_text:
        return ""

    return "\n".join(rows_text)


# ============================================================================
# 语义分块
# ============================================================================


def semantic_chunk_documents(text: str) -> List[LCDocument]:
    """
    使用递归字符分块将文本分割成块

    按段落和字符长度递归分割，保持语义完整性。
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=400,  # 每个块的最大字符数
        chunk_overlap=100,  # 块之间的重叠字符数
        separators=["\n\n", "\n", "。", "！", "？", " ", ""],  # 分割符优先级
    )

    return splitter.create_documents([text])


# ============================================================================
# 文档上传
# ============================================================================


def upload_document(file_content: bytes, filename: str) -> Dict[str, Any]:
    """
    上传并处理文档（支持 docx 和 txt）

    Args:
        file_content: 文件二进制内容
        filename: 文件名

    Returns:
        {"doc_id": str, "chunk_count": int, "chunks": list}
    """
    # 1. 根据文件后缀选择解析方式
    if filename.lower().endswith(".txt"):
        full_text, content_items = parse_txt_document(file_content)
    else:
        full_text, content_items = parse_word_document(file_content)

    if not full_text.strip():
        raise ValueError("文档内容为空")

    # 2. 语义分块
    langchain_docs = semantic_chunk_documents(full_text)

    # 3. 获取 ChromaDB 向量库
    vectorstore = get_chroma_client()

    # 4. 准备数据并存储
    doc_id = str(uuid.uuid4())
    lc_documents = []

    for idx, doc in enumerate(langchain_docs):
        lc_documents.append(
            LCDocument(
                page_content=doc.page_content,
                metadata={
                    "source": filename,
                    "doc_id": doc_id,
                    "chunk_index": idx,
                    "total_chunks": len(langchain_docs),
                },
            )
        )

    vectorstore.add_documents(documents=lc_documents)

    return {
        "doc_id": doc_id,
        "filename": filename,
        "chunk_count": len(lc_documents),
        "chunks": [
            {
                "id": f"{doc_id}_{idx}",
                "content": doc.page_content[:200] + "..."
                if len(doc.page_content) > 200
                else doc.page_content,
            }
            for idx, doc in enumerate(langchain_docs)
        ],
    }


# ============================================================================
# 文档检索 - 基础检索
# ============================================================================


def search_documents(
    query: str, top_k: int = 5, doc_id: Optional[str] = None
) -> Dict[str, Any]:
    """
    搜索文档（基础检索）

    Args:
        query: 查询文本
        top_k: 返回结果数量
        doc_id: 可选，限定在某个文档内搜索

    Returns:
        {"results": [...], "query": str}
    """
    vectorstore = get_chroma_client()

    try:
        docs = vectorstore.similarity_search_with_score(
            query=query,
            k=top_k,
            filter={"doc_id": doc_id} if doc_id else None,
        )
    except Exception:
        return {"results": [], "query": query, "message": "No documents found"}

    # 格式化结果
    formatted_results = []
    for doc, score in docs:
        formatted_results.append(
            {
                "content": doc.page_content,
                "metadata": doc.metadata,
                "distance": 1 - score,  # 转为距离
                "chunk_id": doc.metadata.get("chunk_id", ""),
            }
        )

    return {
        "query": query,
        "results": formatted_results,
        "total": len(formatted_results),
    }


# ============================================================================
# 两阶段检索 + LLM 整合
# ============================================================================


def get_llm_client() -> ChatOpenAI:
    """获取 LLM 客户端（单例）"""
    return ChatOpenAI(
        model=LLM_MODEL,
        api_key=LLM_API_KEY,
        base_url=LLM_API_BASE,
        temperature=0.3,  # 低温度，更准确的回答
    )


def build_context_from_chunks(all_chunks: List[Dict]) -> str:
    """
    将 chunks 构建为 LLM 上下文

    Args:
        all_chunks: 检索到的 chunks 列表

    Returns:
        格式化的上下文字符串
    """
    if not all_chunks:
        return "未找到相关文档内容。"

    # 按文档分组构建上下文
    context_parts = []
    current_source = None

    for chunk in all_chunks:
        source = chunk.get("source", "未知来源")
        content = chunk.get("content", "")

        # 如果换了文档来源，添加分隔
        if source != current_source:
            if current_source is not None:
                context_parts.append("")
            context_parts.append(f"【{source}】")
            current_source = source

        context_parts.append(f"- {content}")

    return "\n".join(context_parts)


def qa_search(query: str, top_k: int = 10) -> Dict[str, Any]:
    """
    问答式搜索（LangGraph 版：查询改写 + 相关性评估 + 最多 1 次重检索）

    工作流：
        rewrite -> retrieve -> grade ->[should_retry]-> rewrite / generate -> END

    对外输出形状与旧实现完全一致，前端无需改动。

    Args:
        query: 用户问题
        top_k: 检索数量（当前 MVP 固定使用 10）

    Returns:
        {
            "answer": str,
            "sources": [...],
            "used_chunks": int,
            "groups": {...},
            "query": str,
            "total_retrieved": int,
            "total_docs": int,
        }
    """
    # 延迟导入，避免循环依赖（rag_graph 在加载时会触发本模块导入）
    from services.rag_graph import rag_graph

    initial_state = {
        "query": query,
        "retry_count": 0,
        "rewritten_query": "",
        "retrieved_chunks": [],
        "relevant_chunks": [],
        "is_relevant": False,
    }

    # LangGraph invoke 是同步的；FastAPI 在线程池中运行，无需 asyncio
    final_state = rag_graph.invoke(initial_state)

    return {
        "answer": final_state.get("answer", ""),
        "sources": final_state.get("sources", []),
        "used_chunks": final_state.get("used_chunks", 0),
        "groups": final_state.get("groups", {}),
        "query": query,
        "total_retrieved": final_state.get("total_retrieved", 0),
        "total_docs": final_state.get("total_docs", 0),
    }


# ============================================================================
# 文档管理
# ============================================================================


def list_documents() -> List[Dict[str, Any]]:
    """列出所有已上传的文档"""
    vectorstore = get_chroma_client()

    try:
        all_docs = vectorstore.get(include=["metadatas"])
    except Exception:
        return []

    if not all_docs or not all_docs.get("ids"):
        return []

    # 按 doc_id 分组
    docs_map = {}
    for i, (doc_id_val, metadata) in enumerate(
        zip(all_docs["ids"], all_docs.get("metadatas", []))
    ):
        if metadata:
            doc_id = metadata.get("doc_id")
            if doc_id and doc_id not in docs_map:
                docs_map[doc_id] = {
                    "doc_id": doc_id,
                    "source": metadata.get("source", "unknown"),
                    "total_chunks": metadata.get("total_chunks", 0),
                    "chunk_ids": [],
                }
            if doc_id:
                docs_map[doc_id]["chunk_ids"].append(doc_id_val)

    return list(docs_map.values())


def delete_document(doc_id: str) -> Dict[str, Any]:
    """删除指定文档的所有块"""
    vectorstore = get_chroma_client()

    try:
        all_docs = vectorstore.get(include=["metadatas"])
    except Exception:
        return {"success": False, "message": "Collection not found"}

    # 获取该文档的所有块 ID
    chunk_ids = []
    for i, (doc_id_val, metadata) in enumerate(
        zip(all_docs["ids"], all_docs.get("metadatas", []))
    ):
        if metadata and metadata.get("doc_id") == doc_id:
            chunk_ids.append(doc_id_val)

    if not chunk_ids:
        return {"success": False, "message": "Document not found"}

    vectorstore.delete(ids=chunk_ids)

    return {"success": True, "deleted_chunks": len(chunk_ids)}
